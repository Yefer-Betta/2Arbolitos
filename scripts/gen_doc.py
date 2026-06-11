import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text=''):
    p = doc.add_paragraph(text)
    return p

def main():
    doc = Document()
    # Título principal
    title = doc.add_heading('Documentación del Proyecto - 2Arbolitos POS', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Índice (se dejará como marcador de posición)
    add_heading(doc, 'Índice', 1)
    doc.add_paragraph('')

    # Estructura del documento según el esquema proporcionado
    sections = [
        ('Lista de Tablas', 1, ''),
        ('Lista de Figuras', 1, ''),
        ('Resumen', 1, 'Resumen ejecutivo del proyecto.'),
        ('Glosario', 1, ''),
        ('Agradecimientos', 1, ''),
        ('1. Planteamiento del Problema', 1, ''),
        ('1.1 Descripción del problema', 2, 'Descripción detallada del problema que se pretende resolver.'),
        ('1.2 Formulación del problema', 2, ''),
        ('1.3 Alcance del proyecto', 2, ''),
        ('Incluye:', 3, ''),
        ('No incluye (fuera del alcance):', 3, ''),
        ('1.4 Objetivos', 2, ''),
        ('1.4.1 Objetivo general', 3, ''),
        ('1.4.2 Objetivos específicos', 3, ''),
        ('1.5 Justificación', 2, ''),
        ('Impacto económico', 3, ''),
        ('Soberanía de datos', 3, ''),
        ('Contexto de hiperinflación venezolana', 3, ''),
        ('Inclusión digital', 3, ''),
        ('Personalización', 3, ''),
        ('Académico (formación SENA)', 3, ''),
        ('2. Análisis de Requerimientos', 1, ''),
        ('2.1 Caracterización de la organización o cliente', 2, ''),
        ('Tipo de cliente', 3, ''),
        ('Perfil típico del restaurante', 3, ''),
        ('Usuarios finales', 3, ''),
        ('Infraestructura típica', 3, ''),
        ('2.2 Levantamiento de requerimientos', 2, ''),
        ('Técnicas utilizadas', 3, ''),
        ('Hallazgos clave', 3, ''),
        ('2.3 Requerimientos funcionales', 2, ''),
        ('2.4 Requerimientos no funcionales', 2, ''),
        ('2.5 Historias de usuario', 2, ''),
        ('HU-01: Tomar pedido (Mesero)', 3, ''),
        ('HU-02: Marcar pedido como listo (Cocinero)', 3, ''),
        ('HU-03: Procesar pago multi‑moneda (Cajero)', 3, ''),
        ('HU-04: Configurar tasas de cambio (Administrador)', 3, ''),
        ('HU-05: Ver cierre de caja multi‑moneda (Cajero o Administrador)', 3, ''),
        ('2.6 Casos de uso', 2, ''),
        ('CU-01: Tomar pedido', 3, ''),
        ('CU-02: Marcar pedido como listo', 3, ''),
        ('CU-03: Cobrar cuenta multi‑moneda', 3, ''),
        ('CU-04: Gestionar menú', 3, ''),
        ('CU-05: Gestionar usuarios', 3, ''),
        ('CU-06: Generar cierre de caja', 3, ''),
        ('CU-07: Sincronizar dispositivos', 3, ''),
        ('2.7 Reglas de negocio', 2, ''),
        ('3. Diseño de la Solución de Software', 1, ''),
        ('3.1 Arquitectura del sistema', 2, ''),
        ('3.2 Diseño de la base de datos', 2, ''),
        ('3.2.1 Modelo entidad‑relación', 3, ''),
        ('3.2.2 Diccionario de datos', 3, ''),
        ('3.3 Diseño de interfaces de usuario', 2, ''),
        ('Paleta de colores', 3, ''),
        ('Tipografía', 3, ''),
        ('Layout principal', 3, ''),
        ('3.4 Diagramas UML', 2, ''),
        ('3.4.1 Diagrama de casos de uso', 3, ''),
        ('3.4.2 Diagrama de clases', 3, ''),
        ('3.4.3 Diagrama de secuencia', 3, ''),
        ('3.4.4 Diagrama de componentes', 3, ''),
        ('3.5 Tecnologías y herramientas', 2, ''),
        ('Frontend', 3, ''),
        ('Backend', 3, ''),
        ('Base de datos', 3, ''),
        ('Empaquetado', 3, ''),
        ('Herramientas de desarrollo', 3, ''),
        ('4. Desarrollo e Implementación', 1, ''),
        ('4.1 GitHub', 2, ''),
        ('Repositorio', 3, ''),
        ('Estrategia de branching', 3, ''),
        ('Estadísticas del repositorio', 3, ''),
        ('Estructura del repositorio', 3, ''),
        ('Commits destacados del proyecto', 3, ''),
        ('Configuración de CI/CD', 3, ''),
        ('5. Conclusiones', 1, ''),
        ('Logros principales', 2, ''),
        ('Aprendizajes del proyecto', 2, ''),
        ('Métricas del proyecto', 2, ''),
        ('6. Recomendaciones', 1, ''),
        ('Técnicas', 2, ''),
        ('Funcionales', 2, ''),
        ('De negocio', 2, ''),
        ('7. Referencias Bibliográficas', 1, ''),
        ('8. Anexos', 1, ''),
        ('Anexo A. Análisis FODA del proyecto', 2, ''),
        ('Anexo B. Matriz de riesgos', 2, ''),
        ('Anexo C. Estimación de costos del proyecto', 2, ''),
        ('Anexo D. Cronograma de actividades (Gantt textual)', 2, ''),
        ('Anexo E. Manual de usuario resumido', 2, ''),
        ('E.1 Mesero', 3, ''),
        ('E.2 Cocinero', 3, ''),
        ('E.3 Cajero', 3, ''),
        ('E.4 Administrador', 3, ''),
        ('Anexo F. Manual de instalación rápido', 2, ''),
        ('F.1 Requisitos previos', 3, ''),
        ('F.2 Pasos de instalación', 3, ''),
        # Wizard parts omitted intentionally
        ('Anexo G. Comparativa con competidores (Tabla 8)', 2, ''),
    ]

    for title, level, placeholder in sections:
        add_heading(doc, title, level)
        if placeholder:
            add_paragraph(doc, placeholder)
        else:
            add_paragraph(doc)

    # Guardar documento
    output_path = Path('Documentación del Proyecto/2Arbolitos_POS_Documento_Proyecto_CORREGIDO.docx')
    doc.save(output_path)
    print(f'Documento creado en {output_path}')

if __name__ == '__main__':
    main()
