import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text=''):
    p = doc.add_paragraph(text)
    return p

def add_centered_paragraph(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return p

def add_mono_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    # Ensure the font is applied to the text run
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rPr.append(rFonts)
    return p

def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Light List Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table

def main():
    doc = Document()
    # ---- Title page (APA) ----
    add_centered_paragraph(doc, 'Servicio Nacional de Aprendizaje (SENA)', size=12, bold=True)
    add_centered_paragraph(doc, 'Regional Bogotá D.C. – Centro de Servicios Financieros', size=12)
    add_centered_paragraph(doc, 'Programa: Análisis y Desarrollo de Software', size=12)
    add_centered_paragraph(doc, 'Evidencia de Proyecto Formativo', size=12)
    doc.add_paragraph('')
    add_centered_paragraph(doc, 'Sistema de Punto de Venta y Gestión Integral para Restaurantes de Zona Fronteriza', size=16, bold=True)
    add_centered_paragraph(doc, 'Soporte multi‑moneda: COP, USD y Bs.', size=14, bold=True)
    doc.add_paragraph('')
    add_centered_paragraph(doc, 'Aprendiz: Yeferson Hotman Pérez Gómez', size=12)
    add_centered_paragraph(doc, 'Instructor: Jesús Antonio Figueroa Guerrero', size=12)
    add_centered_paragraph(doc, 'Ubicación del cliente: El Nula, Estado Apure, Venezuela', size=12)
    add_centered_paragraph(doc, 'Bogotá D.C., Colombia — Junio 2026', size=12)
    add_centered_paragraph(doc, 'Versión 1.0', size=12)
    doc.add_page_break()

    # ---- Abstract (Resumen) ----
    add_heading(doc, 'Resumen', level=1)
    add_paragraph(doc,
        "2Arbolitos POS es un sistema de punto de venta y gestión integral para restaurantes, "
        "diseñado para operar completamente en red local (LAN) sin requerir conexión a internet. "
        "El proyecto se desarrolló como evidencia formativa del programa Análisis y Desarrollo de Software del SENA, "
        "Centro de Servicios Financieros, Regional Bogotá D.C., con el objetivo de resolver las problemáticas de "
        "digitalización que enfrentan los restaurantes pequeños en zonas de frontera colombo‑venezolana, específicamente "
        "en el municipio El Nula, Estado Apure, Venezuela. "
        "El sistema soporta tres monedas (peso colombiano COP, dólar estadounidense USD y bolívar venezolano Bs.) con tasas de "
        "cambio configurables en tiempo real, lo que resulta esencial en una economía bimonetaria como la de la frontera apureña. "
        "La arquitectura cliente‑servidor de tres capas implementa una API REST con Express, persistencia en MySQL mediante el ORM Prisma, "
        "y sincronización en tiempo real entre dispositivos mediante Server‑Sent Events (SSE) con versionado optimista para resolver conflictos de edición. "
        "El frontend se construyó con React 19 y Tailwind CSS, optimizado para interfaces táctiles en tablets y celulares. "
        "El sistema incluye cuatro roles diferenciados (Administrador, Mesero, Cocina y Cajero), módulo de Punto de Venta, vista de mesas en tiempo real, "
        "Kitchen Display System con vista Kanban, gestión de caja con cálculo automático de vueltos, y reportes de cierre diario. "
        "Las pruebas de funcionalidad y sincronización se realizaron con éxito en entorno local, validando el correcto manejo de 5 dispositivos simultáneos, "
        "la resolución de conflictos y la consistencia de datos. La documentación técnica completa incluye 17 archivos en el repositorio de GitHub que cubren análisis, "
        "diseño UML, modelo entidad‑relación, manuales de usuario y de instalación."
    )
    add_paragraph(doc, 'Palabras clave: punto de venta, restaurante, red local, sincronización en tiempo real, multi‑moneda, frontera, Node.js, React, software libre.')
    doc.add_page_break()

    # ---- Índice General (placeholder) ----
    add_heading(doc, 'Índice General', level=1)
    add_paragraph(doc, '(El índice se actualizará automáticamente al abrir el documento en Word).')
    doc.add_page_break()

    # ---- Lista de Tablas ----
    add_heading(doc, 'Lista de Tablas', level=1)
    for t in [
        'Tabla 1. Requerimientos funcionales del sistema.',
        'Tabla 2. Requerimientos no funcionales del sistema.',
        'Tabla 3. Diccionario de datos de las entidades principales.',
        'Tabla 4. Análisis FODA del proyecto 2Arbolitos POS.',
        'Tabla 5. Matriz de riesgos del proyecto.'
    ]:
        add_paragraph(doc, t)
    doc.add_page_break()

    # ---- Lista de Figuras ----
    add_heading(doc, 'Lista de Figuras', level=1)
    for f in [
        'Figura 1. Arquitectura del sistema en 3 capas.',
        'Figura 2. Diagrama de clases del dominio (ASCII).',
        'Figura 3. Diagrama de secuencia del flujo Tomar pedido (ASCII).',
        'Figura 4. Diagrama de componentes del sistema (ASCII).',
        'Figura 5. Paleta de colores corporativa.',
        'Figura 6. Wireframe de la pantalla POS.',
        'Figura 7. Mapa de navegación del sistema.',
        'Figura 8. Topología de red en el restaurante.'
    ]:
        add_paragraph(doc, f)
    doc.add_page_break()

    # ---- Resumen (already in abstract) – we can skip duplicate

    # ---- Glosario ----
    add_heading(doc, 'Glosario', level=1)
    for entry in [
        'API REST: Interfaz de programación de aplicaciones que sigue el estilo arquitectónico REST, usando HTTP para la comunicación entre cliente y servidor.',
        
        'Bcrypt: Algoritmo de hash de passwords basado en Blowfish, usado para almacenar contraseñas de forma segura.',
        'BSSRDE: Bolívar Soberano, moneda oficial de Venezuela desde 2018.',
        'COP: Peso colombiano, moneda oficial de Colombia.',
        'CORS: Cross‑Origin Resource Sharing, mecanismo que permite solicitudes HTTP entre diferentes orígenes.',
        'DIAN: Dirección de Impuestos y Aduanas Nacionales de Colombia, entidad reguladora de la facturación electrónica.',
        'Express: Framework minimalista para Node.js que facilita la creación de servidores HTTP y APIs REST.',
        'KDS: Kitchen Display System, sistema de pantalla de cocina que muestra los pedidos pendientes en tiempo real.',
        'LAN: Local Area Network, red de área local que conecta dispositivos en un espacio reducido como un restaurante.',
        'mDNS: Multicast DNS, protocolo que permite resolver nombres de host en una red local sin servidor DNS central.',
        'MIT: Licencia de software libre permisiva que permite el uso, modificación y distribución del software.',
        'MySQL: Sistema de gestión de bases de datos relacional de código abierto.',

        'POS: Point of Sale, sistema de punto de venta utilizado en comercios para registrar transacciones.',
        'Prisma: ORM (Object‑Relational Mapping) moderno para Node.js que facilita el acceso a bases de datos.',
        'React: Biblioteca de JavaScript para construir interfaces de usuario basada en componentes.',
        'SSE: Server‑Sent Events, tecnología que permite al servidor enviar actualizaciones al cliente en tiempo real sobre HTTP.',
        'Versionado optimista: Técnica de concurrencia que detecta conflictos comparando números de versión antes de aplicar cambios.'
    ]:
        add_paragraph(doc, entry)
    doc.add_page_break()

    # ---- Agradecimientos ----
    add_heading(doc, 'Agradecimientos', level=1)
    for ack in [
        "A Dios, por darme la fortaleza, la sabiduría y la perseverancia necesarias para culminar este proyecto formativo.",
        "A mi familia, por su apoyo incondicional y por ser el motor que impulsa cada uno de mis proyectos personales y académicos.",
        "Al Servicio Nacional de Aprendizaje SENA, por brindarme la oportunidad de formarme profesionalmente en el programa de Análisis y Desarrollo de Software, y por los recursos tecnológicos puestos a mi disposición.",
        "A mi instructor Jesús Antonio Figueroa Guerrero, por su orientación, paciencia y compromiso en la transmisión de conocimiento técnico. Su guía fue fundamental para el éxito de este proyecto.",
        "A los compañeros de formación, con quienes comparto este camino de aprendizaje, por los intercambios de ideas y el apoyo mutuo durante las jornadas de trabajo.",
        "Al restaurante piloto ubicado en El Nula, Estado Apure, Venezuela, por abrir sus puertas y permitirme implementar el sistema en un entorno real, validando así la propuesta tecnológica en un contexto fronterizo colombo‑venezolano.",
        "A la comunidad de desarrolladores de software libre, cuyos aportes en frameworks como React, Node.js y Prisma hicieron posible construir esta solución sin costo de licenciamiento.",
        "A todos, muchas gracias."
    ]:
        add_paragraph(doc, ack)
    doc.add_page_break()

    # ---- 1. Planteamiento del Problema ----
    add_heading(doc, '1. Planteamiento del Problema', level=1)
    add_heading(doc, '1.1 Descripción del problema', level=2)
    add_paragraph(doc,
        "El restaurante piloto beneficiario de este proyecto se ubica en la localidad de El Nula, capital del Municipio Páez del Estado Apure, "
        "en la República Bolivariana de Venezuela, en una zona de frontera con el Departamento de Arauca, Colombia. Esta región presenta "
        "características socioeconómicas particulares que dificultan la adopción de tecnologías de gestión empresarial disponibles en el mercado."
    )
    add_paragraph(doc,
        "La mayoría de las soluciones de Punto de Venta (POS) existentes en el mercado son sistemas SaaS (Software como Servicio) con dependencia "
        "crítica de conexión a internet, tales como Alegra POS, Siigo POS o Square. En El Nula y municipios vecinos, la conexión a internet es "
        "intermitente, costosa y en muchos casos inexistente, lo que hace inviable operar un sistema cien por ciento en la nube. Adicionalmente, "
        "estas plataformas comercializan suscripciones mensuales que oscilan entre $200.000 y $500.000 COP, un costo significativo para pequeños negocios familiares."
    )
    add_paragraph(doc,
        "Otro factor determinante es la naturaleza bimonetaria (y en la práctica tri‑monetaria) de la economía local. En El Nula circulan simultáneamente el peso colombiano (COP), "
        "el dólar estadounidense (USD) como moneda de ahorro, y el bolívar venezolano (Bs.) en sus distintas denominaciones (Bs.S, Bs.D, Bs.Soberano). Los clientes del restaurante "
        "pagan indistintamente en cualquiera de las tres monedas, y el negocio necesita registrar las ventas y calcular vueltos considerando la conversión entre ellas a tasas actualizadas en tiempo real."
    )
    add_paragraph(doc,
        "La consecuencia directa es que el restaurante sigue utilizando métodos manuales (papel, calculadora, cuadernos) para registrar pedidos, calcular cuentas, controlar caja y emitir cierres diarios. "
        "Esto genera errores de cálculo, pérdida de información, demoras en la atención al cliente y ausencia de indicadores para la toma de decisiones."
    )
    add_heading(doc, '1.2 Formulación del problema', level=2)
    add_paragraph(doc,
        "¿Cómo diseñar e implementar un sistema de Punto de Venta y gestión integral para restaurantes pequeños que opere completamente en una red local (LAN) sin requerir "
        "conexión a internet, que soporte múltiples monedas (COP, USD y bolívar venezolano) con tasas de cambio configurables en tiempo real, que sea accesible económicamente para un "
        "negocio familiar en zona de frontera, y que permita al negocio ser dueño de sus propios datos?"
    )
    add_heading(doc, '1.3 Alcance del proyecto', level=2)
    add_paragraph(doc, "Incluye:")
    for inc in [
        "Punto de Venta táctil para meseros con interfaz optimizada para tablets y celulares.",
        "Vista de Mesas en tiempo real con mapa visual y estados (libre, ocupada, cuenta pedida, pagando).",
        "Kitchen Display System (KDS) para cocina con tablero Kanban (pendientes, en preparación, listos).",
        "Gestión de caja con soporte para tres monedas: COP, USD y bolívar venezolano (Bs.).",
        "Cálculo automático de vueltos multi‑moneda con tasas de conversión configurables en tiempo real.",
        "Reportes de cierre diario con desglose por método de pago y por moneda.",
        "Acceso desde celulares y tablets vía Wi‑Fi local.",
        "Sincronización en tiempo real entre dispositivos vía Server‑Sent Events (SSE).",
        "Cuatro roles diferenciados: Administrador, Mesero, Cocina y Cajero.",
        "Código abierto bajo licencia MIT, sin costo de licenciamiento."
    ]:
        add_paragraph(doc, "● " + inc)
    add_paragraph(doc, "No incluye (fuera del alcance):")
    for exc in [
        "Facturación electrónica DIAN (requiere integración con proveedor certificado colombiano).",
        "Integración con datáfonos físicos (solo registro manual del pago con tarjeta).",
        "Módulo de delivery o domicilios.",
        "Inventario avanzado con control de stock y mermas.",
        "Contabilidad formal y nómina.",
        "App móvil nativa para iOS/Android (solo web responsive)."
    ]:
        add_paragraph(doc, "● " + exc)
    add_heading(doc, '1.4 Objetivos', level=2)
    add_heading(doc, '1.4.1 Objetivo general', level=3)
    add_paragraph(doc,
        "Desarrollar un sistema POS integral para restaurantes pequeños de zona fronteriza que funcione en red local sin internet, con soporte multi‑moneda "
        "(COP, USD y bolívar venezolano), interfaz táctil, sincronización en tiempo real, y despliegue mediante Docker."
    )
    add_heading(doc, '1.4.2 Objetivos específicos', level=3)
    for obj in [
        "Diseñar e implementar una arquitectura cliente‑servidor de 3 capas con API REST y eventos en tiempo real (SSE) para sincronización entre dispositivos.",
        "Desarrollar un módulo de Punto de Venta táctil optimizado para tablets y celulares con manejo de mesas y pedidos.",
        "Implementar un Kitchen Display System (KDS) con vista Kanban de pedidos pendientes, en preparación y listos.",
        "Crear un sistema de caja con soporte para tres monedas (COP, USD, Bs.) y cálculo automático de vueltos."
    ]:
        add_paragraph(doc, "● " + obj)
    add_heading(doc, '1.5 Justificación', level=2)
    add_paragraph(doc, "La realización de este proyecto se justifica desde cinco dimensiones:")
    add_paragraph(doc, "Impacto económico: Los restaurantes pequeños en Colombia invierten entre $2.400.000 y $6.000.000 COP al año en suscripciones a POS comerciales. Una solución local con licencia MIT representa un ahorro significativo que se reinvierte directamente en el negocio. Para el restaurante piloto en El Nula, donde la economía está dolarizada y los ingresos son en moneda venezolana devaluada, evitar una suscripción mensual en pesos colombianos es una ventaja competitiva.")
    add_paragraph(doc, "Soberanía de datos: Al funcionar offline‑first, los datos nunca salen del restaurante. No hay dependencia de proveedores externos ni riesgo de pérdida de información por cierre de la empresa SaaS o por cambios en sus políticas de precios. En zonas de frontera con conectividad limitada, esta característica es crítica.")
    add_paragraph(doc, "Contexto de hiperinflación venezolana: La economía venezolana ha enfrentado hiperinflación desde 2016, lo que ha generado una dolarización informal de la economía. Un sistema que maneja simultáneamente bolívar (Bs.), peso colombiano (COP) y dólar estadounidense (USD) refleja fielmente la realidad comercial de El Nula y permite al negocio tomar decisiones informadas sobre qué moneda retener como reserva de valor.")
    add_paragraph(doc, "Inclusión digital: Zonas con internet limitado o nulo pueden acceder a tecnología de gestión moderna mediante este sistema, democratizando el acceso a herramientas digitales en el sector gastronómico de frontera.")
    add_paragraph(doc, "Personalización: El código es abierto (licencia MIT) y modificable. El restaurante o futuros desarrolladores pueden adaptarlo a necesidades específicas, agregar funcionalidades o cambiar la lógica de negocio sin depender de terceros.")
    add_paragraph(doc, "Académico (formación SENA): El proyecto integra conocimientos de ingeniería de software, bases de datos, redes, diseño UX/UI, desarrollo web full‑stack y empaquetado de aplicaciones, sirviendo como evidencia formativa integral del programa Análisis y Desarrollo de Software del SENA.")
    doc.add_page_break()

    # ---- 2. Análisis de Requerimientos ----
    add_heading(doc, '2. Análisis de Requerimientos', level=1)
    add_heading(doc, '2.1 Caracterización de la organización o cliente', level=2)
    add_heading(doc, 'Tipo de cliente', level=3)
    add_paragraph(doc, "Restaurantes pequeños y medianos (5 a 20 dispositivos) ubicados en zonas urbanas y periurbanas, con énfasis particular en zonas de frontera colombo‑venezolana como El Nula, Estado Apure.")
    add_heading(doc, 'Perfil típico del restaurante', level=3)
    for p in [
        "3 a 15 mesas.",
        "1 a 5 meseros.",
        "1 a 2 cocineros.",
        "1 cajero o administrador.",
        "1 dueño o gerente (quien ejerce múltiples roles).",
        "Operación de 10 a 14 horas por día.",
        "Facturación diaria entre $1.000.000 y $10.000.000 COP (o su equivalente en USD/Bs.)."
    ]:
        add_paragraph(doc, "● " + p)
    add_heading(doc, 'Usuarios finales', level=3)
    add_paragraph(doc, "Administrador: dueño o gerente. Acceso total al sistema, gestiona menú, usuarios, configuración y cierres de caja.")
    add_paragraph(doc, "Mesero: toma pedidos, gestiona mesas. Interfaz táctil rápida, operable con una sola mano.")
    add_paragraph(doc, "Cocinero: ve pedidos en pantalla de cocina. Marca pedidos como en preparación y como listos.")
    add_paragraph(doc, "Cajero: procesa pagos en cualquier moneda, realiza cierre de caja diario.")
    add_heading(doc, 'Infraestructura típica', level=3)
    for i in [
        "1 PC servidor (Windows o Linux) con MySQL 8 y la aplicación servidor.",
        "1 a 3 tablets para meseros (Android o iPad).",
        "1 a 2 pantallas adicionales (cocina y caja).",
        "Red Wi‑Fi local con router (no requiere internet).",
        "Conexión eléctrica con planta de respaldo (UPS) opcional."
    ]:
        add_paragraph(doc, "● " + i)
    add_heading(doc, '2.2 Levantamiento de requerimientos', level=2)
    add_heading(doc, 'Técnicas utilizadas', level=3)
    for t in [
        "Entrevistas semiestructuradas con el dueño y personal del restaurante piloto en El Nula, Estado Apure.",
        "Observación directa de operaciones durante 2 jornadas completas de trabajo.",
        "Análisis de sistemas POS comerciales (Alegra, Siigo, Square) para comparativa.",
        "Revisión de la documentación interna de operaciones del restaurante.",
        "Consultas a la comunidad de restaurantes fronterizos para validar funcionalidades requeridas."
    ]:
        add_paragraph(doc, "● " + t)
    add_heading(doc, 'Hallazgos clave', level=3)
    for h in [
        "El proceso de toma de pedido es el más crítico y el principal punto de fricción.",
        "La comunicación mesero‑cocina se pierde cuando hay ruido en el local.",
        "El cierre de caja manual consume entre 30 y 45 minutos al final del día.",
        "El cálculo manual de vueltos en distintas monedas genera errores frecuentes.",
        "La conversión COP a USD y viceversa, así como a bolívar, debe ser automática.",
        "La red Wi‑Fi del restaurante suele ser inestable en horas pico."
    ]:
        add_paragraph(doc, "● " + h)
    add_heading(doc, '2.3 Requerimientos funcionales', level=2)
    add_paragraph(doc, "Tabla 1. Requerimientos funcionales del sistema.")
    rf_rows = [
        ("RF-01", "Crear, editar y eliminar productos del menú con categorías y precios.", "Alta"),
        ("RF-02", "Gestionar el estado de las mesas (libre, ocupada, cuenta pedida, pagando).", "Alta"),
        ("RF-03", "Tomar pedidos asociados a una mesa con cantidades y notas especiales.", "Alta"),
        ("RF-04", "Mostrar los pedidos pendientes en tiempo real en la pantalla de cocina.", "Alta"),
        ("RF-05", "Marcar pedidos como en preparación y como listos desde la pantalla de cocina.", "Alta"),
        ("RF-06", "Cobrar pedidos con múltiples métodos de pago (efectivo, tarjeta, mixto).", "Alta"),
        ("RF-07", "Soportar conversión automática entre COP, USD y bolívar (Bs.) con tasas configurables.", "Alta"),
        ("RF-08", "Generar un ticket imprimible o por correo al cerrar la cuenta.", "Media"),
        ("RF-09", "Generar reporte de cierre de caja diario con totales por método de pago y por moneda.", "Media"),
        ("RF-10", "Permitir que 5 o más dispositivos trabajen simultáneamente sobre la misma base de datos.", "Alta"),
        ("RF-11", "Notificar a todos los dispositivos cuando hay cambios en mesas o pedidos vía SSE.", "Alta"),
        ("RF-12", "Generar código QR para acceso desde celulares vía Wi‑Fi local.", "Media"),
        ("RF-13", "Crear, editar y gestionar usuarios con distintos roles (Admin, Mesero, Cocina, Cajero).", "Alta"),
        ("RF-14", "Permitir hacer backup de los datos en formato JSON.", "Baja"),
        ("RF-15", "Configurar datos del negocio (nombre, logo, NIT) y tasa de cambio multi‑moneda.", "Media"),
        ("RF-16", "Calcular vueltos automáticamente respetando la moneda en que se efectúa el pago.", "Alta"),
        ("RF-17", "Permitir pagos mixtos (parte en efectivo, parte con tarjeta).", "Media")
    ]
    add_table(doc, ["ID", "Requerimiento", "Prioridad"], rf_rows)
    add_heading(doc, '2.4 Requerimientos no funcionales', level=2)
    add_paragraph(doc, "Tabla 2. Requerimientos no funcionales del sistema.")
    rnf_rows = [
        ("RNF-01", "El sistema debe responder en menos de 500 ms en operaciones locales.", "Rendimiento"),
        ("RNF-02", "Soportar al menos 20 dispositivos conectados simultáneamente.", "Escalabilidad"),
        ("RNF-03", "Funcionar sin conexión a internet (offline‑first).", "Disponibilidad"),
        ("RNF-04", "Interfaz táctil operable con dedos (botones >= 44 px de área).", "Usabilidad"),
        ("RNF-05", "Instalable mediante Docker Compose (un solo comando).", "Usabilidad"),
        ("RNF-06", "Datos persistentes en MySQL con respaldo automático diario.", "Confiabilidad"),
        ("RNF-07", "Manejar conflictos de edición con versionado optimista.", "Consistencia"),
        ("RNF-08", "Cifrado de contraseñas con bcrypt y HTTPS en producción.", "Seguridad"),
        ("RNF-09", "Compatible con Windows 10+, macOS 12+, Ubuntu 22.04+.", "Portabilidad"),
        ("RNF-10", "Código siguiendo convenciones ESLint y Prettier.", "Mantenibilidad"),
        ("RNF-11", "Sincronización entre dispositivos en menos de 2 segundos.", "Rendimiento"),
        ("RNF-12", "Interfaz adaptable a distintos tamaños de pantalla (responsive).", "Usabilidad")
    ]
    add_table(doc, ["ID", "Requerimiento", "Categoría"], rnf_rows)
    add_heading(doc, '2.5 Historias de usuario', level=2)
    for hu in [
        ("HU‑01: Tomar pedido (Mesero)",
         "Como mesero, quiero registrar un pedido desde mi tablet tocando los productos, para enviarlo directamente a cocina sin tener que ir a la barra.",
         ["Puedo seleccionar la mesa de un mapa visual.",
          "Veo el menú organizado por categorías.",
          "Agrego productos con cantidad y notas especiales.",
          "Confirmo el pedido y aparece en cocina en menos de 2 segundos."]),
        ("HU‑02: Marcar pedido como listo (Cocinero)",
         "Como cocinero, quiero ver los pedidos en una pantalla grande y marcarlos como en preparación y luego listo, para que el mesero sepa cuándo recoger.",
         ["Veo los pedidos en columnas: Pendientes, En preparación, Listos.",
          "Toco un pedido para cambiar su estado.",
          "El mesero recibe una notificación cuando marco listo."]),
        ("HU‑03: Procesar pago multi‑moneda (Cajero)",
         "Como cajero, quiero cobrar la cuenta de una mesa aceptando pagos en COP, USD o bolívar, para cerrar la venta y emitir el ticket al cliente sin importar la moneda que use.",
         ["Veo el detalle de la cuenta con subtotal, impuestos y total.",
          "Puedo seleccionar la moneda en que recibo el pago.",
          "Ingreso el monto recibido y veo el vuelto calculado automáticamente.",
          "El sistema convierte entre monedas usando la tasa actual.",
          "Se genera un ticket con el detalle de la compra."]),
        ("HU‑04: Configurar tasas de cambio (Administrador)",
         "Como administrador, quiero configurar la tasa de cambio COP/USD y COP/Bs. desde la interfaz, para reflejar la realidad del mercado en tiempo real.",
         ["Puedo modificar la tasa de cambio en cualquier momento.",
          "El cambio se refleja inmediatamente en todas las pantallas.",
          "El sistema guarda un historial de tasas para auditoría."]),
        ("HU‑05: Ver cierre de caja multi‑moneda (Cajero o Administrador)",
         "Como cajero, quiero ver un reporte del cierre de caja del día con totales por método de pago y por moneda (COP, USD, Bs.), para cuadrar la caja con la realidad.",
         ["Veo el total de ventas del día desglosado por moneda.",
          "Veo el detalle por efectivo COP, efectivo USD, efectivo Bs. y tarjeta.",
          "Puedo comparar el efectivo esperado contra el contado.",
          "Puedo exportar el reporte a PDF."])
    ]:
        add_heading(doc, hu[0], level=3)
        add_paragraph(doc, hu[1])
        for c in hu[2]:
            add_paragraph(doc, "● " + c)
    add_heading(doc, '2.6 Casos de uso', level=2)
    for cu in [
        ("CU‑01: Tomar pedido", "Mesero", "El usuario está autenticado y tiene una mesa seleccionada.",
         ["El mesero selecciona una mesa del mapa.",
          "El sistema muestra el menú con categorías.",
          "El mesero agrega productos con cantidad.",
          "El mesero agrega notas especiales si es necesario.",
          "El mesero confirma el pedido.",
          "El sistema envía el pedido a cocina.",
          "El sistema notifica a todos los dispositivos conectados."],
         "El pedido aparece en la pantalla de cocina."),
        ("CU‑02: Marcar pedido como listo", "Cocinero", "Hay pedidos pendientes en cocina.",
         ["El cocinero ve los pedidos en columnas.",
          "Toca un pedido pendiente y el sistema lo mueve a En preparación.",
          "El cocinero prepara el pedido.",
          "Toca el pedido cuando termina y el sistema lo mueve a Listos.",
          "El sistema notifica al mesero asignado."], "El pedido queda marcado como listo."),
        ("CU‑03: Cobrar cuenta multi‑moneda", "Cajero", "El cliente pide la cuenta.",
         ["El cajero selecciona la mesa y el sistema muestra el detalle.",
          "El cajero selecciona la moneda en que se efectúa el pago.",
          "Si es efectivo, ingresa el monto recibido y el sistema calcula el vuelto.",
          "El sistema convierte el pago a la moneda base (COP) para registro.",
          "El sistema registra el pago y libera la mesa."], "La mesa queda libre y la venta registrada."),
        ("CU‑04: Gestionar menú", "Administrador", "El usuario es administrador.",
         ["Accede al módulo de menú.",
          "Agrega, edita o elimina productos con precio, categoría y disponibilidad.",
          "Los cambios se reflejan en todos los dispositivos en tiempo real."], "El menú queda actualizado en todos los dispositivos."),
        ("CU‑05: Gestionar usuarios", "Administrador", "El usuario es administrador.",
         ["Accede al módulo de usuarios.",
          "Crea un nuevo usuario con rol y contraseña temporal.",
          "El nuevo usuario puede iniciar sesión inmediatamente."], "El usuario queda creado y activo."),
        ("CU‑06: Generar cierre de caja", "Cajero o Administrador", "Es el final del día.",
         ["Accede al cierre de caja y el sistema calcula los totales por moneda.",
          "Cuenta el efectivo real por moneda y el sistema lo compara con el esperado.",
          "El sistema genera el reporte final y bloquea nuevos cobros."], "El cierre queda registrado y la caja cerrada."),
        ("CU‑07: Sincronizar dispositivos", "Sistema (automático)", "Hay dispositivos conectados.",
         ["Un dispositivo realiza un cambio y el servidor lo recibe.",
          "El servidor notifica a los demás dispositivos vía SSE.",
          "Los demás dispositivos actualizan su vista.",
          "Si hay conflicto, se aplica versionado optimista."], "Los dispositivos quedan sincronizados.")
    ]:
        add_heading(doc, cu[0], level=3)
        add_paragraph(doc, f"Actor: {cu[1]}. Precondición: {cu[2]}")
        for step in cu[3]:
            add_paragraph(doc, "● " + step)
        add_paragraph(doc, f"Postcondición: {cu[4]}")
    add_heading(doc, '2.7 Reglas de negocio', level=2)
    for rn in [
        "RN‑01: Un pedido solo puede ser enviado a cocina cuando la mesa está en estado Ocupada o Cuenta pedida.",
        "RN‑02: Un producto solo puede ser eliminado del menú si no está incluido en pedidos activos.",
        "RN‑03: El vuelto se calcula siempre en la moneda con que se paga; no se mezclan monedas en el vuelto.",
        "RN‑04: La tasa de cambio COP/USD, COP/Bs. y USD/Bs. es configurable solo por el administrador y aplica a todas las operaciones del día.",
        "RN‑05: Un usuario no puede eliminarse si tiene operaciones registradas en el día.",
        "RN‑06: El cierre de caja solo puede hacerse una vez por día y solo por el cajero o administrador.",
        "RN‑07: Las mesas tienen un máximo de 8 comensales registrados para efectos de control.",
        "RN‑08: Los pedidos enviados a cocina no pueden modificarse, solo anularse con autorización del administrador.",
        "RN‑09: El sistema mantiene un log de auditoría de todas las operaciones críticas: pagos, cierres y cambios de menú.",
        "RN‑10: La sesión de un usuario expira después de 8 horas de inactividad por seguridad.",
        "RN‑11: La moneda base para el registro contable interno es el peso colombiano (COP); los pagos en USD o Bs. se convierten a la tasa vigente al momento del cobro.",
        "RN‑12: El sistema no procesa pagos en otras monedas distintas a COP, USD o Bs."
    ]:
        add_paragraph(doc, rn)
    doc.add_page_break()

    # ---- 3. Diseño de la Solución de Software ----
    add_heading(doc, '3. Diseño de la Solución de Software', level=1)
    add_heading(doc, '3.1 Arquitectura del sistema', level=2)
    add_paragraph(doc, "Figura 1. Arquitectura del sistema en 3 capas.")
    add_mono_paragraph(doc, "+-----------------------------------------------------------+")
    add_mono_paragraph(doc, "|  Capa 1: Presentación (Frontend React 19 + Vite)          |")
    add_mono_paragraph(doc, "|  Componentes: POS, VistaMesas, Cocina, Caja, Configuración|")
    add_mono_paragraph(doc, "+-----------------------------+-----------------------------+")
    add_mono_paragraph(doc, "                              | HTTP REST + SSE")
    add_mono_paragraph(doc, "+-----------------------------v-----------------------------+")
    add_mono_paragraph(doc, "|  Capa 2: Lógica de Negocio (Backend Node.js + Express)    |")
    add_mono_paragraph(doc, "|  Módulos: Auth, Mesas, Pedidos, Productos, Cierres        |")
    add_mono_paragraph(doc, "|  ORM: Prisma 5  |  Autenticación: JWT + bcrypt            |")
    add_mono_paragraph(doc, "+-----------------------------+-----------------------------+")
    add_mono_paragraph(doc, "                              | Prisma Client (SQL)")
    add_mono_paragraph(doc, "+-----------------------------v-----------------------------+")
    add_mono_paragraph(doc, "|  Capa 3: Datos (Persistencia)                              |")
    add_mono_paragraph(doc, "|  MySQL 8.0+ con 11 modelos definidos en Prisma            |")
    add_mono_paragraph(doc, "+-----------------------------------------------------------+")
    add_paragraph(doc, "Patrones aplicados:")
    for p in [
        "MVC (Modelo‑Vista‑Controlador) en el backend.",
        "Repository mediante Prisma ORM.",
        "Observer mediante Server‑Sent Events (SSE) para sincronización en tiempo real.",
        "Optimistic Locking con version counter para evitar conflictos de edición.",
        "Context Provider en React para state management global."
    ]:
        add_paragraph(doc, "● " + p)
    add_paragraph(doc, "Despliegue físico:")
    for d in [
        "Servidor: 1 PC con Node.js 20 LTS + MySQL 8 en la red local del restaurante.",
        "Clientes: (escritorio) en cocina y caja; navegador web en tablets/celulares de meseros.",
        "Comunicación: HTTP en LAN (puerto 3002), broadcast SSE para eventos."
    ]:
        add_paragraph(doc, "● " + d)
    add_heading(doc, '3.2 Diseño de la base de datos', level=2)
    add_heading(doc, '3.2.1 Modelo entidad‑relación', level=3)
    for entity in [
        "User: usuarios del sistema (admin, mesero, cocina, cajero).",
        "Product: productos del menú.",
        "Category: categorías del menú.",
        "Table: mesas del restaurante.",
        "Order: pedidos de una mesa.",
        "OrderItem: líneas de pedido con cantidad y notas.",
        "TableState: estado actual de cada mesa con versionado optimista.",
        "Closure: cierres de caja.",
        "Expense: gastos operativos.",
        "Setting: configuración del sistema clave‑valor.",
        "AuditLog: registro de auditoría."
    ]:
        add_paragraph(doc, "● " + entity)
    add_paragraph(doc, "Relaciones principales:")
    for rel in [
        "Category 1 a N Product.",
        "Table 1 a 1 TableState.",
        "Table 1 a N Order.",
        "Order 1 a N OrderItem.",
        "Order N a 1 User (mesero asignado).",
        "Product 1 a N OrderItem."
    ]:
        add_paragraph(doc, "● " + rel)
    add_heading(doc, '3.2.2 Diccionario de datos', level=3)
    add_paragraph(doc, "Tabla 3. Diccionario de datos de las entidades principales.")
    dict_rows = [
        ("User", "id", "INT", "PK, AUTO_INCREMENT", "ID único del usuario"),
        ("User", "username", "VARCHAR(50)", "UNIQUE, NOT NULL", "Nombre de usuario"),
        ("User", "password", "VARCHAR(255)", "NOT NULL", "Hash bcrypt"),
        ("User", "role", "ENUM", "NOT NULL", "admin, mesero, cocina, cajero"),
        ("User", "active", "BOOLEAN", "DEFAULT TRUE", "Usuario activo"),
        ("Product", "id", "INT", "PK, AUTO_INCREMENT", "ID único del producto"),
        ("Product", "name", "VARCHAR(100)", "NOT NULL", "Nombre del producto"),
        ("Product", "price", "DECIMAL(10,2)", "NOT NULL", "Precio en COP (moneda base)"),
        ("Product", "categoryId", "INT", "FK a Category", "Categoría del producto"),
        ("Product", "available", "BOOLEAN", "DEFAULT TRUE", "Producto disponible"),
        ("Table", "id", "INT", "PK, AUTO_INCREMENT", "ID único de la mesa"),
        ("Table", "number", "INT", "UNIQUE, NOT NULL", "Número de mesa"),
        ("Table", "capacity", "INT", "DEFAULT 4", "Capacidad de comensales"),
        ("TableState", "id", "INT", "PK", "ID de la mesa"),
        ("TableState", "items", "JSON", "NOT NULL", "Lista de items de la mesa"),
        ("TableState", "version", "INT", "DEFAULT 0", "Contador para versionado optimista"),
        ("TableState", "updatedAt", "DATETIME", "", "Última actualización"),
        ("Order", "id", "INT", "PK, AUTO_INCREMENT", "ID único del pedido"),
        ("Order", "tableId", "INT", "FK a Table", "Mesa asociada"),
        ("Order", "status", "ENUM", "NOT NULL", "pendiente, en_preparacion, listo, entregado, pagado"),
        ("Order", "total", "DECIMAL(10,2)", "NOT NULL", "Total del pedido en COP"),
        ("Order", "createdAt", "DATETIME", "DEFAULT NOW()", "Fecha de creación"),
        ("OrderItem", "id", "INT", "PK, AUTO_INCREMENT", "ID único del item"),
        ("OrderItem", "orderId", "INT", "FK a Order", "Pedido al que pertenece"),
        ("OrderItem", "productId", "INT", "FK a Product", "Producto vendido"),
        ("OrderItem", "quantity", "INT", "NOT NULL", "Cantidad"),
        ("OrderItem", "unitPrice", "DECIMAL(10,2)", "NOT NULL", "Precio unitario en COP"),
        ("OrderItem", "notes", "TEXT", "", "Notas especiales (sin cebolla, término medio, etc.)"),
        ("Closure", "id", "INT", "PK, AUTO_INCREMENT", "ID único del cierre"),
        ("Closure", "date", "DATE", "UNIQUE", "Fecha del cierre"),
        ("Closure", "totalSales", "DECIMAL(10,2)", "", "Total ventas del día en COP"),
        ("Closure", "totalCashCop", "DECIMAL(10,2)", "", "Efectivo recibido en COP"),
        ("Closure", "totalCashUsd", "DECIMAL(10,2)", "", "Efectivo recibido en USD"),
        ("Closure", "totalCashBs", "DECIMAL(10,2)", "", "Efectivo recibido en Bs."),
        ("Closure", "totalCard", "DECIMAL(10,2)", "", "Pagos con tarjeta"),
        ("Closure", "expectedCash", "DECIMAL(10,2)", "", "Efectivo esperado"),
        ("Closure", "actualCash", "DECIMAL(10,2)", "", "Efectivo contado por el cajero"),
        ("Setting", "key", "VARCHAR(50)", "PK", "Clave de configuración"),
        ("Setting", "value", "TEXT", "", "Valor (tasas de cambio, nombre del negocio, etc.)"),
        ("AuditLog", "id", "INT", "PK, AUTO_INCREMENT", "ID único del log"),
        ("AuditLog", "userId", "INT", "FK a User", "Usuario que realizó la acción"),
        ("AuditLog", "action", "VARCHAR(50)", "NOT NULL", "Tipo de acción"),
        ("AuditLog", "details", "JSON", "", "Detalles adicionales"),
        ("AuditLog", "createdAt", "DATETIME", "DEFAULT NOW()", "Fecha y hora")
    ]
    add_table(doc, ["Tabla", "Campo", "Tipo", "Restricciones", "Descripción"], dict_rows)
    add_heading(doc, '3.3 Diseño de interfaces de usuario', level=2)
    add_paragraph(doc, "Figura 5. Paleta de colores corporativa.")
    add_paragraph(doc, "Paleta de colores:")
    for c in [
        "primary: #1A4D2E (verde oscuro) – Botones principales, header y elementos destacados.",
        "primary-light: #2D6A42 (verde medio) – Estados hover, énfasis secundario.",
        "secondary: #D4A373 (dorado) – Acentos, reloj, elementos decorativos.",
        "background: #F9F7F2 (crema) – Fondo general de la aplicación.",
        "surface: #F0EBE0 (beige claro) – Tarjetas, paneles y modales.",
        "error: #C0392B (rojo) – Errores, alertas y botones destructivos.",
        "success: #27AE60 (verde) – Confirmaciones y estados exitosos."
    ]:
        add_paragraph(doc, "● " + c)
    add_paragraph(doc, "Tipografía:")
    for tip in [
        "Sans‑serif del sistema: ‑apple‑system, BlinkMacSystemFont, Segoe UI, Roboto.",
        "Títulos: 24‑32 px, peso 700, color primary.",
        "Texto cuerpo: 14‑16 px, peso 400, color negro.",
        "Botones: 16‑18 px, peso 600, color blanco sobre fondo primary."
    ]:
        add_paragraph(doc, "● " + tip)
    add_paragraph(doc, "Layout principal:")
    for lt in [
        "Sidebar fija a la izquierda (240 px) con logo, navegación y reloj en tiempo real.",
        "Área de contenido principal a la derecha.",
        "Header con reloj digital, estado de conexión Wi‑Fi y datos del usuario.",
        "Responsive: en móvil el sidebar se oculta y aparece menú hamburguesa."
    ]:
        add_paragraph(doc, "● " + lt)
    add_paragraph(doc, "Figura 6. Wireframe de la pantalla POS.")
    for w in [
        "Mitad izquierda: panel de categorías y productos organizados en grid táctil.",
        "Mitad derecha: carrito de la mesa actual con totales y cálculo multi‑moneda.",
        "Barra superior: número de mesa, total acumulado, botones de acción."
    ]:
        add_paragraph(doc, "● " + w)
    add_paragraph(doc, "Figura 7. Mapa de navegación del sistema.")
    for n in [
        "Login → Selección de rol → Pantalla principal según rol.",
        "Administrador: Mesas, Pedidos, Cocina, Caja, Configuración, Reportes, Cierre.",
        "Mesero: Mapa de Mesas → Selección Mesa → POS → Confirmar pedido.",
        "Cocinero: Pantalla KDS (Kanban de pedidos).",
        "Cajero: Selección mesa → Detalle cuenta → Cobro multi‑moneda → Ticket."
    ]:
        add_paragraph(doc, "● " + n)
    add_heading(doc, '3.4 Diagramas UML', level=2)
    add_heading(doc, '3.4.1 Diagrama de casos de uso', level=3)
    for actor in ["Mesero", "Cocinero", "Cajero", "Administrador", "Sistema (automático)"]:
        add_paragraph(doc, "● " + actor)
    add_paragraph(doc, "Casos de uso principales (ver sección 2.6):")
    for c in ["CU‑01: Tomar pedido.", "CU‑02: Marcar pedido como listo.", "CU‑03: Cobrar cuenta multi‑moneda.", "CU‑04: Gestionar menú.", "CU‑05: Gestionar usuarios.", "CU‑06: Generar cierre de caja.", "CU‑07: Sincronizar dispositivos."]:
        add_paragraph(doc, "● " + c)
    add_heading(doc, '3.4.2 Diagrama de clases', level=3)
    add_paragraph(doc, "Figura 2. Diagrama de clases del dominio (ASCII).")
    add_mono_paragraph(doc, "+---------------------+         +---------------------+")
    add_mono_paragraph(doc, "|      User           |         |     Category        |")
    add_mono_paragraph(doc, "+---------------------+         +---------------------+")
    add_mono_paragraph(doc, "| -id: int            |         | -id: int            |")
    add_mono_paragraph(doc, "| -username: string   |         | -name: string       |")
    add_mono_paragraph(doc, "| -password: string   |         | -description: string|")
    add_mono_paragraph(doc, "| -role: enum         |         +----------+----------+")
    add_mono_paragraph(doc, "| -active: boolean    |                    | 1")
    add_mono_paragraph(doc, "+---------------------+                    |")
    add_mono_paragraph(doc, "                                           | *")
    add_mono_paragraph(doc, "                                  +---------v--------+")
    add_mono_paragraph(doc, "                                  |     Product     |")
    add_mono_paragraph(doc, "                                  +-----------------+")
    add_mono_paragraph(doc, "                                  | -id: int        |")
    add_mono_paragraph(doc, "                                  | -name: string   |")
    add_mono_paragraph(doc, "                                  | -price: decimal |")
    add_mono_paragraph(doc, "                                  | -available: bool|")
    add_mono_paragraph(doc, "                                  +--------+--------+")
    add_mono_paragraph(doc, "                                           | *")
    add_mono_paragraph(doc, "+---------------------+          * +-------v--------+")
    add_mono_paragraph(doc, "|      Table          |<-------------+  OrderItem   |")
    add_mono_paragraph(doc, "+---------------------+            +---------------+")
    add_mono_paragraph(doc, "| -id: int            |            | -id: int      |")
    add_mono_paragraph(doc, "| -number: int        |            | -quantity: int|")
    add_mono_paragraph(doc, "| -capacity: int      |            | -unitPrice    |")
    add_mono_paragraph(doc, "+----------+----------+            | -notes: string|")
    add_mono_paragraph(doc, "           | 1                     +-------+-------+")
    add_mono_paragraph(doc, "           |                             | *")
    add_mono_paragraph(doc, "           | 1                           |")
    add_mono_paragraph(doc, "           |                   +---------v-------+")
    add_mono_paragraph(doc, "           |                   |     Order       |")
    add_mono_paragraph(doc, "           |                   +-----------------+")
    add_mono_paragraph(doc, "           +------------------>| -id: int        |")
    add_mono_paragraph(doc, "                               | -status: enum   |")
    add_mono_paragraph(doc, "                               | -total: decimal |")
    add_mono_paragraph(doc, "                               | -createdAt: date|")
    add_mono_paragraph(doc, "                               +-----------------+")
    add_heading(doc, '3.4.3 Diagrama de secuencia', level=3)
    add_paragraph(doc, "Figura 3. Diagrama de secuencia del flujo Tomar pedido (ASCII).")
    add_mono_paragraph(doc, "Mesero    Frontend (POS)     Backend (API)     Cocina (KDS)    MySQL")
    add_mono_paragraph(doc, "  |             |                  |                 |              |")
    add_mono_paragraph(doc, "  | 1. Seleccionar mesa           |                 |              |")
    add_mono_paragraph(doc, "  +------------>|                  |                 |              |")
    add_mono_paragraph(doc, "  |             |                  |                 |              |")
    add_mono_paragraph(doc, "  | 2. Agregar productos          |                 |              |")
    add_mono_paragraph(doc, "  +------------>|                  |                 |              |")
    add_mono_paragraph(doc, "  |             |                  |                 |              |")
    add_mono_paragraph(doc, "  | 3. Confirmar pedido           |                 |              |")
    add_mono_paragraph(doc, "  +------------>|                  |                 |              |")
    add_mono_paragraph(doc, "  |             | 4. POST /api/orders                 |              |")
    add_mono_paragraph(doc, "  |             +----------------->|                 |              |")
    add_mono_paragraph(doc, "  |             |                  | 5. INSERT order |              |")
    add_mono_paragraph(doc, "  |             |                  +-----------------+------------->|")
    add_mono_paragraph(doc, "  |             |                  |                 |              |")
    add_mono_paragraph(doc, "  |             |                  | 6. 200 OK       |              |")
    add_mono_paragraph(doc, "  |             |<-----------------+                 |              |")
    add_mono_paragraph(doc, "  |             |                  | 7. SSE notify   |              |")
    add_mono_paragraph(doc, "  |             |                  +---------------->|              |")
    add_mono_paragraph(doc, "  | 8. Pedido confirmado          |                 |              |")
    add_mono_paragraph(doc, "  |<------------+                  |                 | Actualizar UI|")
    add_mono_paragraph(doc, "  |             |                  |                 +--------------+")
    add_heading(doc, '3.4.4 Diagrama de componentes', level=3)
    add_paragraph(doc, "Figura 4. Diagrama de componentes del sistema (ASCII).")
    add_mono_paragraph(doc, "+------------------------------------------------------------+")
    add_mono_paragraph(doc, "|                    Frontend (React 19)                     |")
    add_mono_paragraph(doc, "|  +--------+ +--------+ +--------+ +--------+ +--------+    |")
    add_mono_paragraph(doc, "|  |  POS   | | Mesas  | |Cocina  | | Caja   | |Ajustes |    |")
    add_mono_paragraph(doc, "|  +----+---+ +----+---+ +----+---+ +----+---+ +----+---+    |")
    add_mono_paragraph(doc, "|       |          |          |          |          |         |")
    add_mono_paragraph(doc, "|  +----v----------v----------v----------v----------v-----+   |")
    add_mono_paragraph(doc, "|  |       Context Providers (Orders, Menu, User, ...)   |   |")
    add_mono_paragraph(doc, "|  +----------------------+------------------------------+   |")
    add_mono_paragraph(doc, "+----------------------+-----------------------------------+")
    add_mono_paragraph(doc, "                       | fetch + EventSource")
    add_mono_paragraph(doc, "+----------------------v-----------------------------------+")
    add_mono_paragraph(doc, "|                    Backend (Node.js 20)                   |")
    add_mono_paragraph(doc, "|  +-------+ +-------+ +-------+ +-------+ +-------+        |")
    add_mono_paragraph(doc, "|  | Auth  | |Orders | |Product| | Tables | |Settings|       |")
    add_mono_paragraph(doc, "|  +---+---+ +---+---+ +---+---+ +---+---+ +---+---+        |")
    add_mono_paragraph(doc, "|      |         |         |         |         |             |")
    add_mono_paragraph(doc, "|  +---v---------v---------v---------v---------v-----+       |")
    add_mono_paragraph(doc, "|  |                Prisma ORM 5                       |      |")
    add_mono_paragraph(doc, "|  +----------------------+----------------------------+      |")
    add_mono_paragraph(doc, "+----------------------+-----------------------------------+")
    add_mono_paragraph(doc, "                       | SQL")
    add_mono_paragraph(doc, "              +--------v--------+")
    add_mono_paragraph(doc, "              |     MySQL 8    |")
    add_mono_paragraph(doc, "              +----------------+")
    add_heading(doc, '3.5 Tecnologías y herramientas', level=2)
    add_paragraph(doc, "Frontend:")
    for f in [
        "React 19: biblioteca de UI con componentes funcionales y hooks.",
        "Vite 5: empaquetador y dev server ultrarrápido.",
        "Tailwind CSS 3: framework de CSS utility‑first para diseño responsive.",
        "React Context API: state management global.",
        "Lucide React: iconografía SVG consistente."
    ]:
        add_paragraph(doc, "● " + f)
    add_paragraph(doc, "Backend:")
    for b in [
        "Node.js 20 LTS: runtime JavaScript.",
        "Express 4: framework HTTP minimalista.",
        "Prisma 5: ORM moderno con type‑safety para MySQL.",
        "bcrypt: hash seguro de contraseñas.",
        "jsonwebtoken: autenticación con JWT.",
        "cors: soporte CORS para clientes web.",
        "bonjour-service: mDNS para discovery de dispositivos en LAN."
    ]:
        add_paragraph(doc, "● " + b)
    add_paragraph(doc, "Base de datos:")
    for db in [
        "MySQL 8.0 o superior: RDBMS robusto y maduro.",
        "XAMPP: distribución portable para entorno de desarrollo."
    ]:
        add_paragraph(doc, "● " + db)
    add_paragraph(doc, "Herramientas de desarrollo:")
    for h in [
        "Vite PWA: Service Worker para modo offline web.",
        "ESLint y Prettier: linting y formateo de código.",
        "Git y GitHub: control de versiones y colaboración.",
        "PowerShell y Batch: scripts de build y despliegue en Windows.",
        "Sharp: procesamiento de imágenes (variantes del logo).",
        "Visual Studio Code: editor de código principal."
    ]:
        add_paragraph(doc, "● " + h)
    doc.add_page_break()

    # ---- 4. Desarrollo e Implementación ----
    add_heading(doc, '4. Desarrollo e Implementación', level=1)
    add_heading(doc, '4.1 GitHub', level=2)
    add_paragraph(doc, "Repositorio")
    add_paragraph(doc, "URL: https://github.com/Yefer-Betta/2Arbolitos")
    add_paragraph(doc, "Tipo: público. Licencia: MIT.")
    add_paragraph(doc, "Estrategia de branching")
    for b in [
        "Rama principal: main (siempre estable, desplegable).",
        "Features desarrolladas en commits directos a main por ser un proyecto pequeño y personal.",
        "Tags de versión: v1.0.0 al cierre del proyecto."
    ]:
        add_paragraph(doc, "● " + b)
    add_paragraph(doc, "Estadísticas del repositorio")
    for e in [
        "Aproximadamente 80 commits desde el inicio del proyecto.",
        "17 archivos de documentación técnica en la carpeta docs/.",
        "README profesional con 7 badges informativos.",
        "Workflow de CI configurado en .github/workflows/docker-build.yml."
    ]:
        add_paragraph(doc, "● " + e)
    add_paragraph(doc, "Estructura del repositorio")
    add_mono_paragraph(doc, "2Arbolitos/")
    add_mono_paragraph(doc, "|-- src/                    # Frontend React 19")
    add_mono_paragraph(doc, "|-- server/                 # Backend Node.js + Prisma")
    add_mono_paragraph(doc, "|-- electron/               # Wrapper de escritorio")
    add_mono_paragraph(doc, "|-- scripts/                # CLI multiplataforma")
    add_mono_paragraph(doc, "|-- docs/                   # 17 documentos técnicos")
    add_mono_paragraph(doc, "|-- public/                 # Assets estáticos")
    add_mono_paragraph(doc, "|-- build/                  # Recursos para electron-builder")
    add_mono_paragraph(doc, "|-- logo/                   # Logo fuente")
    add_mono_paragraph(doc, "|-- package.json            # Builds generados (gitignored)")
    add_mono_paragraph(doc, "|-- README.md")
    add_mono_paragraph(doc, "+-- .gitignore")
    doc.add_page_break()

    # ---- 5. Conclusiones ----
    add_heading(doc, '5. Conclusiones', level=1)
    add_paragraph(doc,
        "El desarrollo de 2Arbolitos POS cumplió con los objetivos planteados al inicio del proyecto, entregando un sistema funcional, "
        "probado y documentado que resuelve una necesidad real del mercado de restaurantes en zonas de frontera colombo‑venezolana, "
        "en particular para el restaurante piloto ubicado en El Nula, Estado Apure, Venezuela."
    )
    add_heading(doc, 'Logros principales', level=2)
    for l in [
        "Sistema POS completo y funcional con cuatro roles diferenciados, gestión de mesas, pedidos en tiempo real, cocina KDS y caja multi‑moneda (COP, USD, Bs.), operando exitosamente en un entorno de pruebas con 5 o más dispositivos simultáneos.",
        "Sincronización en tiempo real robusta implementada con Server‑Sent Events (SSE) y versionado optimista, eliminando la pérdida de datos y los conflictos de edición que son comunes en sistemas similares.",
        "Soporte completo para tres monedas (peso colombiano, dólar estadounidense y bolívar venezolano) con tasas de cambio configurables en tiempo real, satisfaciendo una necesidad específica del mercado fronterizo.",
        "Arquitectura cliente‑servidor limpia con separación de responsabilidades (3 capas), aplicando patrones reconocidos como MVC, Repository y Observer, lo que facilita el mantenimiento y la extensión futura.",
        "Documentación técnica y académica completa de 17 archivos (aproximadamente 4 500 líneas) que cubre análisis de requerimientos, diseño UML, modelo entidad‑relación, manual de usuario y manual técnico.",
        "Funcionamiento 100 % offline en red local, lo que representa una ventaja competitiva significativa frente a soluciones SaaS y permite su uso en zonas con conectividad limitada como El Nula."
    ]:
        add_paragraph(doc, "● " + l)
    add_heading(doc, 'Aprendizajes del proyecto', level=2)
    for a in [
        "La implementación de sincronización en tiempo real requiere un diseño cuidadoso del versionado para evitar pérdida de datos bajo condiciones de red inestable.",
        "El versionado optimista con contador entero resultó más limpio que mantener el contador embebido en JSON.",
        "La decisión de no firmar digitalmente el ejecutable en esta versión (auto‑firmado) es aceptable para distribución interna o académica, pero requerirá firma EV para distribución comercial.",
        "El contexto multi‑moneda exige diseñar el modelo de datos pensando en una moneda base (COP) para los registros contables, con campos adicionales para almacenar el pago original en la moneda en que se recibió."
    ]:
        add_paragraph(doc, "● " + a)
    doc.add_page_break()

    # ---- 7. Referencias Bibliográficas (APA) ----
    add_heading(doc, '7. Referencias Bibliográficas', level=1)
    refs = [
        "Servicio Nacional de Aprendizaje (SENA). (2026). Evidencia de proyecto formativo: 2Arbolitos POS. Bogotá D.C., Colombia.",
        "React. (2024). React – A JavaScript library for building user interfaces. https://reactjs.org/",
        "Node.js. (2024). Node.js – JavaScript runtime. https://nodejs.org/",
        "Prisma. (2024). Prisma – Next‑generation ORM. https://www.prisma.io/",
        "Docker. (2024). Docker – Enterprise Container Platform. https://www.docker.com/",

        "MySQL. (2024). MySQL – The world’s most popular open‑source database. https://www.mysql.com/"
    ]
    for r in refs:
        add_paragraph(doc, r)
    doc.add_page_break()

    # ---- 8. Anexos ----
    add_heading(doc, '8. Anexos', level=1)
    add_heading(doc, 'Anexo A. Análisis FODA del proyecto', level=2)
    add_paragraph(doc, "El siguiente análisis FODA (Fortalezas, Oportunidades, Debilidades, Amenazas) sintetiza la situación estratégica del sistema 2Arbolitos POS:")
    add_paragraph(doc, "Tabla 4. Análisis FODA del proyecto 2Arbolitos POS.")
    foda_rows = [
        ("Fortaleza", "Funcionamiento offline‑first", "No requiere conexión a internet; opera 100 % en red local."),
        ("Fortaleza", "Soporte multi‑moneda", "Maneja COP, USD y bolívar venezolano con tasas configurables."),
        ("Fortaleza", "Sincronización en tiempo real", "SSE con versionado optimista garantiza consistencia de datos."),
        ("Fortaleza", "Código abierto MIT", "Personalizable, auditable y sin costo de licenciamiento."),
        ("Fortaleza", "4 roles diferenciados", "Permisos granulares por tipo de usuario."),
        ("Oportunidad", "Mercado fronterizo desatendido", "Pocas soluciones POS nativas en español para frontera colombo‑venezolana."),
        ("Oportunidad", "Hiperinflación venezolana", "Impulsa adopción de sistemas digitales con manejo multi‑moneda."),
        ("Oportunidad", "Crecimiento gastronómico", "Sector en expansión en ambos países."),
        ("Oportunidad", "Formación SENA", "Posibilidad de escalar el proyecto a más aprendices."),
        ("Debilidad", "Sin firma digital EV", "Muestra advertencia de Windows SmartScreen al instalar."),
        ("Debilidad", "No incluye facturación electrónica DIAN", "Requiere integración futura con proveedor certificado."),
        ("Debilidad", "Sin delivery ni domicilios", "Módulo no implementado en esta versión."),
        ("Debilidad", "Sin inventario avanzado", "Control de stock no contemplado."),
        ("Debilidad", "App móvil no es nativa", "Solo web responsive; no hay app iOS/Android."),
        ("Amenaza", "Competencia de SaaS", "Alegra, Siigo, Square con mayor inversión en marketing."),
        ("Amenaza", "Crisis económica venezolana", "Reduce capacidad de pago de los restaurantes."),
        ("Amenaza", "Piratería de software", "Facilidad de copia no licenciada."),
        ("Amenaza", "Inestabilidad eléctrica y de red", "Frecuentes en zonas rurales y de frontera."),
        ("Amenaza", "Cambio de versiones de MySQL o Node.js", "Puede romper compatibilidad a futuro.")
    ]
    add_table(doc, ["Categoría", "Aspecto", "Detalle"], foda_rows)
    add_heading(doc, 'Anexo B. Matriz de riesgos', level=2)
    add_paragraph(doc, "La siguiente matriz identifica los principales riesgos del proyecto, su probabilidad, impacto y estrategia de mitigación:")
    add_paragraph(doc, "Tabla 5. Matriz de riesgos del proyecto.")
    riesgo_rows = [
        ("R‑01", "Pérdida de datos por fallo de hardware", "Baja", "Crítico", "Backup automático diario + RAID 1 + UPS"),
        ("R‑02", "Internet inestable en El Nula", "Alta", "Alto", "Diseño offline‑first con sincronización SSE"),
        ("R‑03", "Hiperinflación afecta precio del sistema", "Alta", "Medio", "Multi‑moneda con tasa configurable en tiempo real"),
        ("R‑04", "Cambio de versión de Node.js o MySQL", "Baja", "Alto", "Prisma abstrae la BD; pin de versiones en package.json"),
        ("R‑05", "Falta de capacitación del personal", "Media", "Alto", "Asistente gráfico + manual de usuario + soporte remoto"),
        ("R‑06", "Piratería del software", "Alta", "Bajo", "Licencia MIT: libre por diseño, no hay pérdida económica"),
        ("R‑07", "Resistencia al cambio del personal", "Media", "Medio", "Capacitación previa + interfaz intuitiva"),
        ("R‑08", "Falla del servidor MySQL", "Baja", "Crítico", "Monitoreo + reinicio automático + backup"),
        ("R‑09", "Sincronización incorrecta entre dispositivos", "Baja", "Alto", "Versionado optimista + conflict‑merge + tests"),
        ("R‑10", "Robo de equipos (tablets)", "Media", "Alto", "Marcado de activos + seguro + backup de datos en servidor")
    ]
    add_table(doc, ["ID", "Riesgo", "Probabilidad", "Impacto", "Mitigación"], riesgo_rows)
    add_heading(doc, 'Anexo E. Manual de usuario resumido', level=2)
    add_paragraph(doc, "A continuación se presenta un manual de usuario resumido para los cuatro roles del sistema:")
    add_heading(doc, 'E.1 Mesero', level=3)
    for m in [
        "Inicia sesión con tu usuario y contraseña.",
        "En la pantalla principal, toca el icono “Mesas” para ver el mapa del restaurante.",
        "Selecciona una mesa libre tocando sobre ella.",
        "Aparece el menú organizado por categorías. Agrega productos tocando sobre ellos.",
        "Para agregar notas (ej. “sin cebolla”), toca el producto en el carrito y escribe la nota.",
        "Confirma el pedido tocando “Enviar a cocina”. El pedido aparecerá automáticamente en la pantalla de cocina."
    ]:
        add_paragraph(doc, "● " + m)
    add_heading(doc, 'E.2 Cocinero', level=3)
    for c in [
        "Inicia sesión con tu usuario.",
        "La pantalla muestra el tablero Kanban: Pendientes, En preparación, Listos.",
        "Cuando inicies a preparar un pedido, tócalo y se mueve a “En preparación”.",
        "Cuando termines, tócalo de nuevo y se mueve a “Listos”. El mesero será notificado."
    ]:
        add_paragraph(doc, "● " + c)
    add_heading(doc, 'E.3 Cajero', level=3)
    for c in [
        "Inicia sesión con tu usuario.",
        "Toca el icono “Caja” y selecciona la mesa a cobrar.",
        "Aparece el detalle de la cuenta. Selecciona el método de pago (efectivo, tarjeta o mixto).",
        "Si es efectivo, selecciona la moneda (COP, USD o Bs.) e ingresa el monto recibido.",
        "El sistema calcula el vuelto automáticamente. Confirma el cobro.",
        "El sistema genera un ticket con el detalle. Puedes imprimirlo o enviarlo por correo.",
        "Al final del día, realiza el “Cierre de caja” desde el menú principal."
    ]:
        add_paragraph(doc, "● " + c)
    add_heading(doc, 'E.4 Administrador', level=3)
    for a in [
        "Inicia sesión con tu usuario administrador.",
        "Desde el menú “Configuración” puedes: gestionar el menú, usuarios, tasa de cambio multi‑moneda, datos del negocio, backup y restauración.",
        "Desde “Reportes” puedes ver el historial de cierres de caja y exportar a PDF.",
        "Desde “Servidor” puedes ver el estado de sincronización y configurar el inicio automático con Windows."
    ]:
        add_paragraph(doc, "● " + a)

    # Save document
    output_path = Path('Documentación del Proyecto/2Arbolitos_POS_Documento_Proyecto_APA.docx')
    doc.save(output_path)
    print(f'Documento APA creado en {output_path}')

if __name__ == '__main__':
    main()
